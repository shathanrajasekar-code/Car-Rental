import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner margins (padding) for a table cell in twentieths of a point (dxa)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = docx.oxml.OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = docx.oxml.OxmlElement(m)
        node.set(docx.oxml.ns.qn('w:w'), str(val))
        node.set(docx.oxml.ns.qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()
    
    # 1. Setup Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Setup Header/Footer
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("Car Rental Administration System | IET Report")
        hrun.font.name = 'Times New Roman'
        hrun.font.size = Pt(8.5)
        hrun.font.italic = True
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("Page ")
        frun.font.name = 'Times New Roman'
        frun.font.size = Pt(9)
    
    # 2. Base font styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    
    # Heading style configurations
    for i in range(1, 4):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.name = 'Times New Roman'
        h_style.font.bold = True
        h_style.font.color.rgb = RGBColor(0, 0, 0)
    
    # Set heading sizes
    doc.styles['Heading 1'].font.size = Pt(18)
    doc.styles['Heading 2'].font.size = Pt(14)
    doc.styles['Heading 3'].font.size = Pt(12)

    def add_para(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        return p

    def add_heading_1(text):
        p = doc.add_heading(text, level=1)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.color.rgb = RGBColor(0, 0, 0)
        return p

    def add_heading_2(text):
        p = doc.add_heading(text, level=2)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.color.rgb = RGBColor(0, 0, 0)
        return p

    def add_heading_3(text):
        p = doc.add_heading(text, level=3)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.color.rgb = RGBColor(0, 0, 0)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        return p

    # ==========================================
    # COVER PAGE
    # ==========================================
    add_para("A PROJECT REPORT ON", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=36, space_after=12)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(24)
    r_title = p_title.add_run("CAR RENTAL ADMINISTRATION SYSTEM\n(BHARATH RENTAL SYSTEM)")
    r_title.bold = True
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(24)
    
    add_para("Submitted in partial fulfillment of the requirements for the award of the degree of", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para("BACHELOR OF SCIENCE IN COMPUTER SCIENCE", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    
    add_para("Submitted By", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para("DEVI PRIYANKA. S\n(REG NO: 21BCT009)", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)
    
    add_para("Under the Guidance of", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para("Mr. R. Karthik, MCA., M.Phil.\nAssistant Professor, Department of Computer Science", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=48)
    
    # Institution Info
    add_para("DEPARTMENT OF COMPUTER SCIENCE", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para("SRI KRISHNA ARTS AND SCIENCE COLLEGE", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para("(Affiliated to Bharathiar University, Coimbatore)", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para("COIMBATORE — 641008, TAMIL NADU", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para("JULY 2026", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12)
    
    doc.add_page_break()

    # ==========================================
    # DECLARATION
    # ==========================================
    add_heading_1("DECLARATION")
    add_para("", space_before=12)
    
    dec_text = (
        "I, DEVI PRIYANKA. S (Reg. No: 21BCT009), hereby declare that the project report entitled "
        "\"CAR RENTAL ADMINISTRATION SYSTEM (BHARATH RENTAL SYSTEM)\" submitted to Sri Krishna Arts "
        "and Science College, Coimbatore, in partial fulfillment of the requirements for the award of the "
        "degree of Bachelor of Science in Computer Science is a record of original project work done "
        "by me under the guidance of Mr. R. Karthik, MCA., M.Phil., Assistant Professor, Department "
        "of Computer Science, Sri Krishna Arts and Science College, and that it has not formed the basis "
        "for the award of any other degree, diploma, fellowship, or associate-ship of this or any other University."
    )
    add_para(dec_text, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_para("", space_before=24)
    
    # Signature Lines
    p_sig = doc.add_paragraph()
    p_sig.paragraph_format.space_before = Pt(48)
    r_place = p_sig.add_run("Place: Coimbatore\t\t\t\t\tSignature of the Candidate\nDate:\t\t\t\t\t\t(DEVI PRIYANKA. S)")
    r_place.font.name = 'Times New Roman'
    r_place.font.size = Pt(12)
    
    doc.add_page_break()

    # ==========================================
    # CERTIFICATE
    # ==========================================
    add_heading_1("CERTIFICATE")
    add_para("", space_before=12)
    
    cert_text = (
        "This is to certify that the project report entitled \"CAR RENTAL ADMINISTRATION SYSTEM "
        "(BHARATH RENTAL SYSTEM)\" is a bona fide record of work done by DEVI PRIYANKA. S "
        "(Reg. No: 21BCT009) in partial fulfillment of the requirements for the award of the degree "
        "of Bachelor of Science in Computer Science of Sri Krishna Arts and Science College, Coimbatore, "
        "during the academic year 2025 - 2026."
    )
    add_para(cert_text, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_para("", space_before=36)
    
    # Signatures
    p_signers = doc.add_paragraph()
    p_signers.paragraph_format.space_before = Pt(48)
    r_signers = p_signers.add_run(
        "Signature of the Guide\t\t\t\tSignature of the HOD\n"
        "(Mr. R. Karthik)\t\t\t\t\t(Dr. Sunitha C.)"
    )
    r_signers.bold = True
    r_signers.font.name = 'Times New Roman'
    r_signers.font.size = Pt(12)
    
    p_exam = doc.add_paragraph()
    p_exam.paragraph_format.space_before = Pt(48)
    r_exam = p_exam.add_run("Submitted for the Viva-Voce Examination held on ____________________")
    r_exam.font.name = 'Times New Roman'
    r_exam.font.size = Pt(12)
    
    p_examiners = doc.add_paragraph()
    p_examiners.paragraph_format.space_before = Pt(36)
    r_examiners = p_examiners.add_run("Internal Examiner\t\t\t\t\tExternal Examiner")
    r_examiners.bold = True
    r_examiners.font.name = 'Times New Roman'
    r_examiners.font.size = Pt(12)
    
    doc.add_page_break()

    # ==========================================
    # ACKNOWLEDGEMENT
    # ==========================================
    add_heading_1("ACKNOWLEDGEMENT")
    add_para("", space_before=12)
    
    ack_1 = (
        "First, I express my deep gratitude to our Almighty God for giving me the strength and wisdom "
        "to complete this project work successfully."
    )
    add_para(ack_1, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    ack_2 = (
        "I express my sincere thanks to the Management of Sri Krishna Arts and Science College for "
        "providing the state-of-the-art laboratory infrastructure and computing facilities that made "
        "this software project possible."
    )
    add_para(ack_2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    ack_3 = (
        "I express my deepest gratitude to our respected Principal, for their guidance, encouragement, "
        "and administrative support throughout the course of my study."
    )
    add_para(ack_3, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    ack_4 = (
        "I am extremely thankful to Dr. Sunitha C., Head of the Department of Computer Science, for "
        "her valuable suggestions, guidance, and academic oversight which kept me focused and motivated."
    )
    add_para(ack_4, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    ack_5 = (
        "I express my profound gratitude to my project guide, Mr. R. Karthik, MCA., M.Phil., Assistant "
        "Professor, Department of Computer Science, for his invaluable guidance, constant supervision, "
        "scholarly insights, and encouragement during the design, coding, testing, and documentation phases "
        "of this Car Rental Administration System."
    )
    add_para(ack_5, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    ack_6 = (
        "Lastly, I express my warm appreciation to my parents, family members, and friends for their constant "
        "moral support, tolerance, and helpful cooperation throughout my academic journey."
    )
    add_para(ack_6, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    doc.add_page_break()

    # ==========================================
    # TABLE OF CONTENTS
    # ==========================================
    add_heading_1("TABLE OF CONTENTS")
    add_para("", space_before=12)
    
    # Manual Table of Contents built in a clean double-column table
    toc_table = doc.add_table(rows=1, cols=3)
    toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    toc_table.style = 'Table Grid'
    
    hdr_cells = toc_table.rows[0].cells
    hdr_cells[0].text = 'Chapter / Section'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Page No'
    for c in hdr_cells:
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.name = 'Times New Roman'
        
    toc_items = [
        ("—", "Cover Page", "i"),
        ("—", "Declaration", "ii"),
        ("—", "Certificate", "iii"),
        ("—", "Acknowledgement", "iv"),
        ("Chapter 1", "Organization Profile", "1"),
        ("Chapter 2", "Abstract", "4"),
        ("Chapter 3", "Introduction", "6"),
        ("3.1", "Overview of the Project", "6"),
        ("3.2", "Objectives", "8"),
        ("3.3", "Scope of the System", "9"),
        ("3.4", "Problem Statement", "10"),
        ("3.5", "Need for the System", "11"),
        ("Chapter 4", "System Study", "12"),
        ("4.1", "Existing System & Drawbacks", "12"),
        ("4.2", "Proposed System & Advantages", "13"),
        ("4.3", "Requirements Analysis Specs", "15"),
        ("Chapter 5", "System Design & Description", "18"),
        ("5.1", "Data Flow Diagrams (DFD 0, 1, 2)", "18"),
        ("5.2", "Entity Relationship (ER) Diagram", "21"),
        ("5.3", "Database Schema Tables Design", "24"),
        ("5.4", "Input and Output Design", "29"),
        ("5.5", "Module Descriptions", "33"),
        ("Chapter 6", "System Testing", "41"),
        ("6.1", "Testing Methods (Unit, Integration, Black Box)", "41"),
        ("6.2", "Test Case Matrix Table", "44"),
        ("Chapter 7", "Conclusion", "47"),
        ("Chapter 8", "Future Scope", "49"),
        ("Chapter 9", "Bibliography", "51"),
        ("Chapter 10", "Appendix & Time Sheet", "53"),
    ]
    
    for ch_sec, desc, pg in toc_items:
        row_cells = toc_table.add_row().cells
        row_cells[0].text = ch_sec
        row_cells[1].text = desc
        row_cells[2].text = pg
        for cell in row_cells:
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            
    doc.add_page_break()

    # ==========================================
    # CHAPTER 1: ORGANIZATION PROFILE
    # ==========================================
    add_heading_1("CHAPTER 1: ORGANIZATION PROFILE")
    add_para("", space_before=12)
    
    add_para(
        "This system was engineered under the software development structure of a modern web technologies division, "
        "dedicated to producing enterprise-grade solutions for consumer transport automation. The enterprise focuses "
        "on replacing fragmented physical booking processes with secure cloud-native platforms tailored for "
        "dynamic markets. With specialized units in database normalization, full-stack microservices, and client-side "
        "interface design, the organization is committed to deploying software that stands out for both premium aesthetics "
        "and robust performance.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    
    add_heading_2("1.1 Vision")
    add_para(
        "To deliver cutting-edge software solutions that simplify business operations, driving digital transformation "
        "through robust, secure, and user-friendly web platforms that combine analytical monitoring with automated workflows.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    
    add_heading_2("1.2 Mission (Goal)")
    add_para(
        "To architect and implement reliable, secure, and scalable web applications utilizing state-of-the-art "
        "languages and frameworks, ensuring high performance, mathematical data consistency, absolute security, "
        "and continuous customer engagement through premium visual designs.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    
    add_heading_2("1.3 Core Services & Company Philosophy")
    add_para(
        "The software development team adheres to a strict code quality philosophy, which includes detailed database "
        "schema planning, REST API standardization, modular UI componentization, and extensive test coverage. "
        "The organization offers a range of high-end software development services:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    add_bullet("Full Stack Web Development (MERN, React, Express, Node.js)")
    add_bullet("Enterprise Database Architecture & Optimization (Mongoose, MongoDB, SQL)")
    add_bullet("Automated Process Scripting & CLI Tool Design (Python, Batch, Shell)")
    add_bullet("Cloud-Native API Deployment and Microservice Orchestration")
    add_bullet("Automated Software Testing, Quality Assurance, and Security Audits")
    
    add_heading_2("1.4 Technology Stack Utilized")
    add_para(
        "The development division works primarily with modern open-source web tech stacks to build responsive, "
        "non-blocking applications. For this project, the core tech stack includes:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    add_bullet("Frontend: React.js (Vite compiler), Tailwind CSS v3, Recharts, jsPDF")
    add_bullet("Backend API: Node.js (V8 engine), Express.js framework, CORS, JSON Web Token (JWT)")
    add_bullet("Database: MongoDB (Document-oriented NoSQL), Mongoose ORM validation layer")
    add_bullet("Automation Runner: Python 3, Setuptools CLI script wrappers, PowerShell scripts")
    
    add_heading_2("1.5 Software Development Process Model")
    add_para(
        "The organization utilizes the Agile Scrum Methodology for all software development lifecycles. "
        "This ensures iterative progression, quick turnarounds, and robust validation:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    add_bullet("Requirement Gathering & Analysis: Detailing user stories, database entities, and workflows.")
    add_bullet("Sprint Planning & Design: Creating ERDs, DFDs, and prototyping UI components in React.")
    add_bullet("Iterative Implementation: Writing modular code in frontend React and backend Express controllers.")
    add_bullet("Testing & Continuous Integration: Running automated test scripts, lint checks (Oxlint), and code reviews.")
    add_bullet("Deployment: Packaging the workspace with Python CLI wrappers for easy execution and installation.")
    
    doc.add_page_break()

    # ==========================================
    # CHAPTER 2: ABSTRACT
    # ==========================================
    add_heading_1("CHAPTER 2: ABSTRACT")
    add_para("", space_before=12)
    
    add_para(
        "The Car Rental Administration System (Bharath Rental System) is a complete full-stack web application "
        "designed to replace traditional, manual, and error-prone car rental processes with a secure, automated, "
        "and visually premium online portal. Developed using the MongoDB, Express, React, and Node (MERN) stack, "
        "the application separates functions between two main user roles: Customers and Administrators. "
        "The system enforces complex business rules, such as preventing overlapping booking intervals for the same vehicle "
        "and maintaining structured state transitions (Pending -> Confirmed -> Ongoing -> Completed -> Cancelled) "
        "for the vehicle rental lifecycle.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    add_para(
        "Beyond core CRUD operations, the project introduces localized academic features tailored for the Indian "
        "market. These features include an AI FASTag & Toll Auto-Estimator that calculates route-based tolls, an "
        "interactive FASTag Toll Gate Crossing Simulator, SafarLock Speed Warning Alerts that simulate expressway "
        "and city speed limits and log violations, GreenYatra Carbon Offset tracking for environmental contributions, "
        "and the Bharath Yatra Itinerary Planner that auto-recommends pitstops based on selected routes.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    add_para(
        "Security is maintained through JWT-based user session authentication and bcrypt password hashing. "
        "To simplify setup for academic presentations and developer deployment, the application is packaged "
        "with an automated Python CLI tool. This tool handles environment validation, dependency installation, "
        "database seeding, and concurrent server launch. The result is a robust, production-ready project "
        "that demonstrates advanced software design patterns, clean database relationships, and professional front-end design.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    
    doc.add_page_break()

    # ==========================================
    # CHAPTER 3: INTRODUCTION
    # ==========================================
    add_heading_1("CHAPTER 3: INTRODUCTION")
    add_para("", space_before=12)
    
    add_heading_2("3.1 Overview of the Project")
    add_para(
        "The Car Rental Administration System is a comprehensive web portal designed to manage vehicle rental "
        "lifecycles. The system is split into a frontend client built in React (utilizing Tailwind CSS for sleek "
        "glassmorphic UIs and Recharts for analytics) and a backend REST API built in Node.js/Express, communicating "
        "with a MongoDB database. The core purpose is to provide customers with an interactive catalog to browse, filter, "
        "and book vehicles, and to provide admins with a detailed control panel to monitor fleet status, bookings, "
        "user access, and financial performance.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    add_para(
        "The application integrates several custom business logics. For instance, booking dates must be valid, pick-up date "
        "cannot be in the past, and a vehicle cannot be reserved for dates that conflict with existing active bookings. "
        "Billing is computed automatically based on daily prices, security deposits, optional toll pre-payments, and "
        "carbon offset fees. Invoices are generated client-side in PDF format using the jsPDF library, ensuring no overhead "
        "on the server side.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    
    add_heading_2("3.2 Objectives")
    add_para(
        "The primary objectives of the developed system are as follows:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    add_bullet("To automate the manual car rental booking and billing lifecycle to eliminate human scheduling errors.")
    add_bullet("To design a responsive, premium user interface with cascading dropdown filters based on brand and model.")
    add_bullet("To implement a robust state transition engine for bookings: Pending, Confirmed, Ongoing, Completed, and Cancelled.")
    add_bullet("To integrate localized academic widgets: AI FASTag Toll Estimator, SafarLock Speed Alerts, and GreenYatra Carbon Tracker.")
    add_bullet("To build a secure admin panel for inventory CRUD, user blocking, analytical graphs, and CSV data reports.")
    
    add_heading_2("3.3 Scope of the System")
    add_para(
        "The scope of the Car Rental Administration System encompasses all aspects of modern web-based booking solutions. "
        "It includes secure JWT authentication, password hashing, and user profile management. For vehicles, it covers "
        "multi-city fleet management (Chennai, Bengaluru, Mumbai, etc.) with detailed features and rental status states. "
        "For transactions, it automates date-conflict checking, mock credit card checkouts, security deposits, and "
        "PDF invoice generation. The system is designed to run efficiently on standard local systems with MongoDB, Node, "
        "and Python, providing a self-contained environment for demonstrations.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    
    add_heading_2("3.4 Problem Statement")
    add_para(
        "Traditional car rental services in developing markets rely heavily on physical ledgers, phone calls, or disjointed "
        "web forms. This leads to common business problems: double-booking of vehicles, disputes over reservation dates, "
        "lack of real-time availability updates, and manual calculation of billing rates. Furthermore, rental agencies "
        "lack tools to monitor driver behavior (such as speed limit violations), estimate route tolls dynamically, "
        "or offer eco-conscious customers carbon offset options. The lack of a unified admin panel for fleet status "
        "tracking and financial reporting prevents agencies from scaling efficiently.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    
    add_heading_2("3.5 Need for the System")
    add_para(
        "To address these inefficiencies, there is a clear need for a centralized Car Rental Administration System. "
        "By implementing an online portal with an automated booking engine, customers can instantly check vehicle "
        "availability, calculate total trip costs (including tolls and deposits), and make secure mock payments. "
        "The inclusion of tracking widgets (SafarLock Speed monitoring and FASTag crossings) provides an interactive experience, "
        "bridging the gap between software logic and real-world transport operations. For administrators, having a single "
        "dashboard with analytical charts and downloadable reports eliminates administrative overhead and improves "
        "fleet deployment strategies.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    
    doc.add_page_break()

    # ==========================================
    # CHAPTER 4: SYSTEM STUDY
    # ==========================================
    add_heading_1("CHAPTER 4: SYSTEM STUDY")
    add_para("", space_before=12)
    
    add_heading_2("4.1 Existing System & Drawbacks")
    add_para(
        "The existing system in local car rental businesses is heavily manual. Customers must physically visit the office "
        "or make phone calls to inquire about car models, prices, and availability. The receptionist logs these details in "
        "paper registers or simple spreadsheets. When a car is rented, bills are hand-written, adding up daily charges, "
        "mileage fees, and safety deposits manually, which often leads to mathematical errors and customer disputes.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    add_para("The primary drawbacks of the existing manual system are:", bold=True)
    add_bullet("High vulnerability to double-booking errors due to lack of real-time database locks.")
    add_bullet("No automated verification of booking date ranges and availability overlap.")
    add_bullet("Lack of transparency in billing, toll estimations, and security deposit refunds.")
    add_bullet("No mechanism to track speed violations or log route dhabas and pitstops.")
    add_bullet("Time-consuming manual creation of reports, invoices, and fleet performance charts.")
    
    add_heading_2("4.2 Proposed System & Advantages")
    add_para(
        "The proposed Car Rental Administration System (Bharath Rental System) is a cloud-ready MERN stack web application "
        "that fully automates the rental workflow. Customers browse a dynamic catalog, select cars using cascading brand-model "
        "dropdowns, and book reservations. The system automatically performs overlap checks on the database before allowing "
        "a booking. It calculates deposits, estimated tolls, and carbon offsets, providing a transparent breakdown. "
        "Administrators manage the fleet, track bookings, block users, view visual income reports, and export raw CSV records.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    add_para("Key advantages of the proposed system include:", bold=True)
    add_bullet("Automated availability checking that prevents date overlap bookings.")
    add_bullet("Secure authentication utilizing bcrypt hashing and JSON Web Tokens.")
    add_bullet("Interactive academic modules like SafarLock speed triggers and FASTag crossing logs.")
    add_bullet("Dynamic admin control panel with interactive Recharts graphs and CSV exports.")
    add_bullet("Client-side generation of professional PDF invoices, reducing server workload.")
    
    add_heading_2("4.3 Requirements Analysis Specifications")
    
    add_heading_3("4.3.1 Functional Requirements")
    add_bullet("User Management: Registration, Login (JWT-secured), Profile edit, and User blocking by Admins.")
    add_bullet("Vehicle Inventory: CRUD operations by Admin, fleet filtering by Category, Brand, Model, and City location.")
    add_bullet("Booking Engine: Interactive calendar reservation, automatic conflict checks, state machine transitions.")
    add_bullet("Toll & Route Simulator: Route selection (e.g., Delhi-Agra Yamuna Expressway), pre-pay tolls option, interactive toll gate crossing widget.")
    add_bullet("SafarLock Warning: Interactive dashboard speedometer simulating overspeeding (alerts stored in database and shown to Admin).")
    add_bullet("GreenYatra Tracker: Carbon offset contribution tracking per rental checkout.")
    add_bullet("Invoicing & Reports: Direct download of PDF invoices and Admin export of booking/user statistics to CSV.")
    
    add_heading_3("4.3.2 Non-Functional Requirements")
    add_bullet("Security: Passwords must be hashed using bcrypt; REST endpoints must validate JWT authorization headers.")
    add_bullet("Performance: Catalog page load times under 2 seconds; booking conflict validations completed under 500ms.")
    add_bullet("Reliability: Structured Mongoose validations to prevent database pollution with invalid schemas.")
    add_bullet("Usability: Sleek glassmorphic components and intuitive responsive layouts powered by Tailwind CSS.")
    add_bullet("Scalability: Stateless REST APIs that allow scaling the Express server independently of MongoDB.")
    
    add_heading_3("4.3.3 Software Requirements")
    add_bullet("Operating System: Windows 10/11, macOS, or Linux.")
    add_bullet("Database Server: MongoDB Community Server (v5.0 or higher).")
    add_bullet("Runtime Environment: Node.js (v16.0 or higher).")
    add_bullet("Package Manager: npm (v8.0 or higher).")
    add_bullet("Code Compiler/Bundler: Vite (v4.0 or higher).")
    add_bullet("CLI Script Runner: Python (v3.6 or higher) with setuptools.")
    
    add_heading_3("4.3.4 Hardware Requirements")
    add_bullet("Processor: Intel Core i3 / AMD Ryzen 3 or higher.")
    add_bullet("RAM: 4 GB minimum (8 GB recommended for concurrent backend, frontend, and database services).")
    add_bullet("Hard Disk Space: 500 MB free space for project source code and node_modules.")
    add_bullet("Network: Localhost connection (no internet required for local offline executions).")
    
    doc.add_page_break()

    # ==========================================
    # CHAPTER 5: SYSTEM DESIGN
    # ==========================================
    add_heading_1("CHAPTER 5: SYSTEM DESIGN")
    add_para("", space_before=12)
    
    add_heading_2("5.1 Data Flow Diagrams (DFD)")
    add_para(
        "Data Flow Diagrams show how data moves through the Car Rental Administration System, mapping "
        "inputs to outputs across different logical levels.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    
    add_heading_3("5.1.1 DFD Level 0 (Context Diagram)")
    add_para(
        "The Level 0 DFD illustrates the boundary of the system, showing the major external entities (Customer and Admin) "
        "and their primary interactions with the Car Rental System.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    
    # Mermaid representation for DFD Level 0
    add_para("Mermaid Code - DFD Level 0 Context Diagram:", bold=True)
    dfd_0_code = (
        "```mermaid\n"
        "graph TD\n"
        "    Customer((Customer)) -- Register / Book / Pay / Simulator inputs --> System[Car Rental System]\n"
        "    System -- PDF Invoices / Toll status / Dashboards --> Customer\n"
        "    Admin((Administrator)) -- Manage Vehicles / Bookings / Block users --> System\n"
        "    System -- Analytics charts / Speed alerts / CSV exports --> Admin\n"
        "```"
    )
    add_para(dfd_0_code, italic=True, space_after=12)
    
    add_heading_3("5.1.2 DFD Level 1 (Process Breakdown Diagram)")
    add_para(
        "The Level 1 DFD decomposes the system into main logical processes: Authentication, Vehicle Browsing, "
        "Booking Engine, Payment processing, Toll simulation, and Admin management.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    add_para("Mermaid Code - DFD Level 1:", bold=True)
    dfd_1_code = (
        "```mermaid\n"
        "graph TD\n"
        "    C[Customer] --> P1[1.0 Authenticate User]\n"
        "    C --> P2[2.0 Browse & Filter Vehicles]\n"
        "    C --> P3[3.0 Booking & Overlap Engine]\n"
        "    C --> P4[4.0 Simulate Toll & Speed]\n"
        "    \n"
        "    P1 --> D1[(Users Database)]\n"
        "    P2 --> D2[(Vehicles Database)]\n"
        "    P3 --> D3[(Bookings Database)]\n"
        "    P3 --> D4[(Payments Database)]\n"
        "    \n"
        "    A[Admin] --> P5[5.0 Manage Inventory]\n"
        "    A --> P6[6.0 Audit Reports & System Control]\n"
        "    P5 --> D2\n"
        "    P6 --> D3\n"
        "    P6 --> D1\n"
        "```"
    )
    add_para(dfd_1_code, italic=True, space_after=12)

    add_heading_3("5.1.3 DFD Level 2 (Booking & Overlap Sub-process)")
    add_para(
        "Level 2 DFD details the core booking engine process, showing the date-validation checks and how "
        "availability conflicts are handled.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    add_para("Mermaid Code - DFD Level 2 (Booking Overlap):", bold=True)
    dfd_2_code = (
        "```mermaid\n"
        "graph TD\n"
        "    C[Customer] -- Booking Request (Dates) --> V1[5.3.1 Validate Dates Input]\n"
        "    V1 -- Valid Dates --> V2{5.3.2 Query Conflicting Bookings}\n"
        "    V2 -- Conflict Exists --> Err[Return Overlap Error]\n"
        "    V2 -- Available --> V3[5.3.3 Create Pending Booking Record]\n"
        "    V3 --> D3[(Bookings Database)]\n"
        "    V3 -- Trigger payment --> V4[5.3.4 Process Card Mock Checkout]\n"
        "    V4 --> D4[(Payments Database)]\n"
        "    V4 -- Update Status --> V5[5.3.5 Confirm Booking & Lock Vehicle]\n"
        "    V5 --> D2[(Vehicles Database)]\n"
        "```"
    )
    add_para(dfd_2_code, italic=True, space_after=12)

    add_heading_2("5.2 Entity Relationship (ER) Diagram")
    add_para(
        "The Entity-Relationship Diagram defines the logical associations between the primary models: User, Vehicle, "
        "Booking, Payment, and Review.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    add_para("Mermaid Code - ER Diagram:", bold=True)
    er_code = (
        "```mermaid\n"
        "erDiagram\n"
        "    USER ||--o{ BOOKING : places\n"
        "    VEHICLE ||--o{ BOOKING : reserved_in\n"
        "    BOOKING ||--|| PAYMENT : triggers\n"
        "    BOOKING ||--o| REVIEW : receives\n"
        "    USER ||--o{ REVIEW : writes\n"
        "```"
    )
    add_para(er_code, italic=True, space_after=12)

    add_heading_2("5.3 Database Schema Tables Design")
    add_para(
        "The database design consists of five collections in MongoDB managed via Mongoose schemas. "
        "Below are the field properties, primary keys, relationships, and constraints.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )

    # User Table
    add_heading_3("5.3.1 Users Collection Schema")
    t_user = doc.add_table(rows=1, cols=5)
    t_user.style = 'Table Grid'
    hdr = t_user.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = "Field Name", "Data Type", "Constraint", "Key", "Description"
    for c in hdr: c.paragraphs[0].runs[0].bold = True
    
    u_fields = [
        ("_id", "ObjectId", "Required, Auto", "PK", "Unique user identifier"),
        ("name", "String", "Required", "—", "User's full name"),
        ("email", "String", "Required, Unique, Regex", "—", "User's email (login credential)"),
        ("phone", "String", "Required", "—", "User's phone contact"),
        ("password", "String", "Required, Minlen 6", "—", "Hashed password (bcrypt)"),
        ("role", "String", "Enum ('admin', 'customer')", "—", "User authorization role"),
        ("drivingLicense", "String", "Required if customer", "—", "License identifier for verification"),
        ("isBlocked", "Boolean", "Default: false", "—", "Indicates if admin blocked this client")
    ]
    for fn, dt, cn, ky, ds in u_fields:
        row = t_user.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text, row[4].text = fn, dt, cn, ky, ds
        for cell in row: set_cell_margins(cell, 60, 60, 100, 100)

    # Vehicle Table
    add_heading_3("5.3.2 Vehicles Collection Schema")
    t_vehicle = doc.add_table(rows=1, cols=5)
    t_vehicle.style = 'Table Grid'
    hdr = t_vehicle.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = "Field Name", "Data Type", "Constraint", "Key", "Description"
    for c in hdr: c.paragraphs[0].runs[0].bold = True
    
    v_fields = [
        ("_id", "ObjectId", "Required, Auto", "PK", "Unique vehicle identifier"),
        ("name", "String", "Required", "—", "Vehicle display name"),
        ("brand", "String", "Required", "—", "Manufacturer brand name"),
        ("model", "String", "Required", "—", "Vehicle model code"),
        ("year", "Number", "Required", "—", "Year of manufacture"),
        ("category", "String", "Enum ('Sedan', 'SUV', ...)", "—", "Body category style"),
        ("transmission", "String", "Enum ('Manual', 'Automatic')", "—", "Gear shift transmission type"),
        ("fuelType", "String", "Enum ('Petrol', 'Diesel', ...)", "—", "Engine fuel supply type"),
        ("seats", "Number", "Required", "—", "Maximum seating capacity"),
        ("dailyPrice", "Number", "Required", "—", "Rental rate charged per day"),
        ("image", "String", "Required", "—", "Relative path to file upload"),
        ("plateNumber", "String", "Required, Unique", "—", "Registration license plate number"),
        ("status", "String", "Enum ('Available', 'Rented', ...)", "—", "Current state status"),
        ("city", "String", "Enum ('Chennai', 'Mumbai', ...)", "—", "Geographic location node"),
        ("features", "Array [String]", "Default: []", "—", "Convenience items lists"),
        ("rules", "Array [String]", "Default: []", "—", "Rental usage restrictions lists")
    ]
    for fn, dt, cn, ky, ds in v_fields:
        row = t_vehicle.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text, row[4].text = fn, dt, cn, ky, ds
        for cell in row: set_cell_margins(cell, 60, 60, 100, 100)

    # Booking Table
    add_heading_3("5.3.3 Bookings Collection Schema")
    t_booking = doc.add_table(rows=1, cols=5)
    t_booking.style = 'Table Grid'
    hdr = t_booking.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = "Field Name", "Data Type", "Constraint", "Key", "Description"
    for c in hdr: c.paragraphs[0].runs[0].bold = True
    
    b_fields = [
        ("_id", "ObjectId", "Required, Auto", "PK", "Unique booking identifier"),
        ("customer", "ObjectId", "Required, Ref: User", "FK", "Reference to the booking user"),
        ("vehicle", "ObjectId", "Required, Ref: Vehicle", "FK", "Reference to the rented vehicle"),
        ("pickupDate", "Date", "Required", "—", "Start date of rental"),
        ("returnDate", "Date", "Required", "—", "End date of rental"),
        ("totalDays", "Number", "Required", "—", "Computed duration of trip"),
        ("totalAmount", "Number", "Required", "—", "Computed checkout sum price"),
        ("status", "String", "Enum ('Pending', 'Confirmed', ...)", "—", "Current workflow stage"),
        ("payment", "ObjectId", "Ref: Payment", "FK", "Reference to associated transaction"),
        ("securityDeposit", "Number", "Default: 5000", "—", "Refundable safety deposit fee"),
        ("tollRoute", "String", "Default: ''", "—", "Selected expressway toll route"),
        ("tollEstimatedAmount", "Number", "Default: 0", "—", "Estimated cost of tolls"),
        ("prepayTolls", "Boolean", "Default: false", "—", "If tolls were pre-paid at checkout"),
        ("tollPaidAmount", "Number", "Default: 0", "—", "Toll charges deducted in simulator"),
        ("tollLogs", "Array [Object]", "plazaName, charge, time", "—", "History log of toll crossings"),
        ("carbonOffset", "Boolean", "Default: false", "—", "If customer chose carbon offsetting"),
        ("overspeedAlertsCount", "Number", "Default: 0", "—", "Logged speed limit violations count")
    ]
    for fn, dt, cn, ky, ds in b_fields:
        row = t_booking.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text, row[4].text = fn, dt, cn, ky, ds
        for cell in row: set_cell_margins(cell, 60, 60, 100, 100)

    # Payment Table
    add_heading_3("5.3.4 Payments Collection Schema")
    t_payment = doc.add_table(rows=1, cols=5)
    t_payment.style = 'Table Grid'
    hdr = t_payment.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = "Field Name", "Data Type", "Constraint", "Key", "Description"
    for c in hdr: c.paragraphs[0].runs[0].bold = True
    
    p_fields = [
        ("_id", "ObjectId", "Required, Auto", "PK", "Unique payment transaction ID"),
        ("booking", "ObjectId", "Required, Ref: Booking", "FK", "Reference to the associated booking"),
        ("amount", "Number", "Required", "—", "Total monetary amount processed"),
        ("status", "String", "Enum ('Pending', 'Success', 'Failed')", "—", "Status of transaction processing"),
        ("transactionId", "String", "Required, Unique", "—", "External gateway mockup reference"),
        ("method", "String", "Default: 'Card'", "—", "Method of checkout payment"),
        ("createdAt", "Date", "Default: Date.now", "—", "Timestamp of checkout transaction")
    ]
    for fn, dt, cn, ky, ds in p_fields:
        row = t_payment.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text, row[4].text = fn, dt, cn, ky, ds
        for cell in row: set_cell_margins(cell, 60, 60, 100, 100)

    # Review Table
    add_heading_3("5.3.5 Reviews Collection Schema")
    t_review = doc.add_table(rows=1, cols=5)
    t_review.style = 'Table Grid'
    hdr = t_review.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = "Field Name", "Data Type", "Constraint", "Key", "Description"
    for c in hdr: c.paragraphs[0].runs[0].bold = True
    
    r_fields = [
        ("_id", "ObjectId", "Required, Auto", "PK", "Unique review identifier"),
        ("customer", "ObjectId", "Required, Ref: User", "FK", "Reference to reviewing customer"),
        ("vehicle", "ObjectId", "Required, Ref: Vehicle", "FK", "Reference to reviewed vehicle"),
        ("booking", "ObjectId", "Required, Ref: Booking", "FK", "Reference to trip booking node"),
        ("rating", "Number", "Required, Min 1, Max 5", "—", "Star rating assigned (1-5)"),
        ("comment", "String", "Required", "—", "Customer's textual feedback"),
        ("createdAt", "Date", "Default: Date.now", "—", "Timestamp review created")
    ]
    for fn, dt, cn, ky, ds in r_fields:
        row = t_review.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text, row[4].text = fn, dt, cn, ky, ds
        for cell in row: set_cell_margins(cell, 60, 60, 100, 100)

    add_heading_3("5.3.6 Database Normalization & Consistency")
    add_para(
        "Though MongoDB is a document-based NoSQL database, relational integrity is maintained at the application level "
        "using Mongoose population hooks. Normalization is strictly followed up to the Third Normal Form (3NF) equivalent:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    add_bullet("1NF (First Normal Form): All database columns store atomic, single-valued attributes. Features and rules are arrays of strings without sub-elements.")
    add_bullet("2NF (Second Normal Form): All non-key fields are fully dependent on the unique primary key identifier (_id). The vehicle table is independent of users, and booking references both using ObjectIds.")
    add_bullet("3NF (Third Normal Form): Transitive dependency is avoided. For example, payment logs contain booking references rather than duplicating user/vehicle details, and users contain roles rather than duplicating billing parameters.")

    add_heading_2("5.4 Input and Output Design")
    
    add_heading_3("5.4.1 Input Design")
    add_para(
        "Input forms are designed to enforce type validation, required fields, and boundary constraints before sending payloads to Express routers:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    add_bullet("Login Form: Email format regex match and password string check.")
    add_bullet("Registration Form: Capture Name, Phone, Email, Password, and Driving License number (required only for customers).")
    add_bullet("Vehicle Entry Form (Admin): Add Vehicle Name, Brand, Model, Year, Category selection, Transmission, Fuel Type, Seats, Daily Price, Plate Number, City location, and Image upload.")
    add_bullet("Booking Calendar Form: Date pickers for pick-up and return dates. Enforces returnDate >= pickupDate >= Today.")
    add_bullet("Payment Form: Mock Card Details (16-digit card number, expiry date, CVV) for payment authentication simulator.")
    add_bullet("Toll & Simulator Widget: Select routes and click buttons to trigger mock toll crossings and speedometer speed limits.")

    add_heading_3("5.4.2 Output Design")
    add_para(
        "Output screens are tailored to provide visual clarity and download capabilities:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    add_bullet("Customer Dashboard: Visual booking tiles, active rentals, and toll billing details.")
    add_bullet("Admin Dashboard Overview: Monthly income chart (Recharts bar chart), count cards for Total Income, Active Bookings, and Total Cars.")
    add_bullet("Booking Confirmation Card: Summary showing dates, duration, total pricing, deposit fees, and offset contributions.")
    add_bullet("PDF Invoice: Client-side generated file with a clean layout containing invoice header, metadata, pricing breakdown table, and payment status.")
    add_bullet("CSV Export: Clean table datasets downloadable directly from the admin panel containing vehicle logs, booking histories, or user accounts.")

    add_heading_2("5.5 Module Descriptions")
    add_para(
        "The application is structured into modular microservice modules that handle specific tasks in the system lifecycle:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )

    add_heading_3("5.5.1 Authentication Module")
    add_para(
        "Purpose: Handles register, login, session validation, and password cryptography.\n"
        "Workflow: Customer registers -> Hashed password saved. Customer logs in -> Express generates signed JWT. JWT is stored in localStorage by Client.\n"
        "Tables Used: Users collection.\n"
        "Validation: Matches password, enforces unique email constraints, verifies token signatures in routing middleware.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )

    add_heading_3("5.5.2 Customer Module")
    add_para(
        "Purpose: Provides catalog browsing, brand-model cascading dropdown filtering, booking history views, review submission, and simulator panels.\n"
        "Workflow: Browses vehicles -> Checks dates -> Checks out via payment -> Simulates tolls -> Submits ratings.\n"
        "Tables Used: Users, Vehicles, Bookings, Payments, Reviews collections.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )

    add_heading_3("5.5.3 Administrator Module")
    add_para(
        "Purpose: Centralized management interface for CRUD vehicles, state changes, blocking users, reviewing speed alerts, and exporting data tables.\n"
        "Workflow: Views Recharts graphs -> Adds cars -> Reviews ongoing bookings -> Blocks user if speed limits are broken.\n"
        "Tables Used: All database collections.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )

    add_heading_3("5.5.4 FASTag Toll Gate & Route Simulator")
    add_para(
        "Purpose: Academic module to simulate real-world route payments.\n"
        "Workflow: Selected route estimates total toll. Crossing simulator logs time and charge on plaza crossing, deducting balance.\n"
        "Tables Used: Bookings collection (tollLogs and tollPaidAmount fields).",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )

    add_heading_3("5.5.5 SafarLock Speed Monitor Module")
    add_para(
        "Purpose: Monitors simulated driving speeds and alerts admins of violations.\n"
        "Workflow: Customer speedometer triggers speed limit checks. Exceeding limit counts violations. Admin panel highlights flagged bookings.\n"
        "Tables Used: Bookings collection (overspeedAlertsCount field).",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )

    doc.add_page_break()

    # ==========================================
    # CHAPTER 6: SYSTEM TESTING
    # ==========================================
    add_heading_1("CHAPTER 6: SYSTEM TESTING")
    add_para("", space_before=12)
    
    add_para(
        "Testing validates the stability, security, functional requirements, and overall user experience "
        "of the Car Rental Administration System before deployment.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    
    add_heading_2("6.1 Testing Methods")
    add_para(
        "A multi-layered testing paradigm was followed to verify all aspects of the application:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    add_bullet("Unit Testing: Validating individual controller functions (e.g. date overlap logic, password hash comparisons, and route-toll estimations).")
    add_bullet("Integration Testing: Testing API endpoints to ensure correct communication between database schemas, express routers, and response models.")
    add_bullet("System Testing: Verifying complete workflows, such as booking checkout triggering database state changes and updating the admin income charts.")
    add_bullet("Black Box Testing: Validating UI inputs, form submissions, and button clicks without checking internal source code logic.")
    add_bullet("White Box Testing: Reviewing schema validations, JWT token expiration, and process tree cleanup paths in the backend CLI.")
    add_bullet("User Acceptance Testing (UAT): Mocking real-world customer usage to ensure the system is intuitive and stable.")

    add_heading_2("6.2 Test Case Matrix Table")
    add_para(
        "The following matrix outlines the key test scenarios executed to validate functional requirements:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    
    t_test = doc.add_table(rows=1, cols=6)
    t_test.style = 'Table Grid'
    hdr = t_test.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text, hdr[5].text = "ID", "Module", "Input / Scenario", "Expected Output", "Actual Output", "Status"
    for c in hdr: c.paragraphs[0].runs[0].bold = True
    
    test_cases = [
        ("TC-01", "Auth", "Register with existing email", "Error message: Email already exists", "Email exists error thrown", "Passed"),
        ("TC-02", "Auth", "Register customer without license", "Validation error: license required", "Validation failed correctly", "Passed"),
        ("TC-03", "Auth", "Login with wrong password", "Error message: Invalid credentials", "Invalid password rejected", "Passed"),
        ("TC-04", "Catalog", "Filter brand 'Tata', model 'Nexon'", "Display only Tata Nexon cars", "Showed Nexon models only", "Passed"),
        ("TC-05", "Booking", "Book overlapping dates on same car", "Overlapping error and block payment", "Rejected overlapping dates", "Passed"),
        ("TC-06", "Booking", "Book pickup date in the past", "Validation error: invalid dates", "Past dates rejected", "Passed"),
        ("TC-07", "Payment", "Submit 15-digit card number", "Invalid card layout warning", "Payment rejected", "Passed"),
        ("TC-08", "Toll Sim", "Click crossing on Mumbai expressway", "Deduct ₹320 and append toll log", "Toll logged and balance updated", "Passed"),
        ("TC-09", "SafarLock", "Exceed simulated speed limit to 130 km/h", "Log violation and increment count", "Violation counter updated", "Passed"),
        ("TC-10", "Admin", "Admin locks a malicious customer account", "Block user from placing bookings", "User blocked, login denied", "Passed"),
        ("TC-11", "Invoice", "Download invoice for completed booking", "Generate PDF containing correct prices", "PDF generated and downloaded", "Passed"),
        ("TC-12", "Report", "Click Export Bookings to CSV", "Download CSV file with booking logs", "CSV downloaded correctly", "Passed")
    ]
    
    for tid, mod, inp, exp, act, stat in test_cases:
        row = t_test.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text, row[4].text, row[5].text = tid, mod, inp, exp, act, stat
        for cell in row: set_cell_margins(cell, 60, 60, 100, 100)
        
    doc.add_page_break()

    # ==========================================
    # CHAPTER 7: CONCLUSION & CHAPTER 8: FUTURE SCOPE
    # ==========================================
    add_heading_1("CHAPTER 7: CONCLUSION")
    add_para("", space_before=12)
    
    add_para(
        "The Car Rental Administration System (Bharath Rental System) has been successfully designed, implemented, "
        "tested, and documented. The portal effectively eliminates the inefficiencies and errors associated with "
        "manual spreadsheets and hand-written billing registers. By utilizing the MERN stack architecture, the system "
        "provides customers with a fast, secure, and visually appealing interface to browse and rent vehicles. "
        "At the same time, it provides administrators with a complete dashboard to control operations, track billing, "
        "and monitor driver behavior.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    add_para(
        "Unique features like the AI FASTag toll estimator, SafarLock speed warnings, and GreenYatra carbon offsetting "
        "address modern needs in the Indian transportation sector. With modular database schema layouts, "
        "strict date validation logic, client-side invoice generation, and packaging scripts, the project stands out "
        "as a comprehensive, academic-level semester project.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    
    add_heading_1("CHAPTER 8: FUTURE SCOPE")
    add_para("", space_before=12)
    
    add_para(
        "While the system currently offers a complete solution for car rentals, there is room for "
        "future functional expansions in commercial settings:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    add_bullet("Online Payment Gateway Integration: Integrating live payment processors like Razorpay, Paytm, or Stripe to replace mock payments.")
    add_bullet("GPS & Geofencing Integration: Installing hardware trackers on vehicles to capture real-time GPS coordinates, automating SafarLock speedometer warnings.")
    add_bullet("Driver Allocation Module: Allowing customers to request a driver, with automated driver booking checks similar to vehicle booking rules.")
    add_bullet("Native Mobile Apps: Developing Flutter or React Native mobile clients to reach more mobile users.")
    add_bullet("AI Recommendation Engine: Suggesting specific vehicles to customers based on their booking history and categories (e.g. Sedan, SUV).")
    add_bullet("Automated SMS & Email Alerts: Setting up services like Twilio or SendGrid to send instant booking confirmations and billing updates.")
    add_bullet("Predictive Maintenance: Using vehicle mileage and age data to predict when cars need servicing, automating status updates to Maintenance.")
    
    doc.add_page_break()

    # ==========================================
    # CHAPTER 9: BIBLIOGRAPHY
    # ==========================================
    add_heading_1("CHAPTER 9: BIBLIOGRAPHY")
    add_para("", space_before=12)
    
    add_heading_2("9.1 Reference Books")
    add_bullet("Flanagan, D. (2020). JavaScript: The Definitive Guide (7th ed.). O'Reilly Media.")
    add_bullet("Chodorow, K. (2013). MongoDB: The Definitive Guide (2nd ed.). O'Reilly Media.")
    add_bullet("Banks, A., & Porcello, E. (2020). Learning React: Modern Patterns for Developing React Apps (2nd ed.). O'Reilly Media.")
    add_bullet("Dayley, B. (2018). Node.js, MongoDB and Angular Web Development (2nd ed.). Addison-Wesley.")
    
    add_heading_2("9.2 Official Documentation & Web References")
    add_bullet("Node.js Official Documentation. https://nodejs.org/docs/")
    add_bullet("React Official Documentation. https://react.dev/")
    add_bullet("MongoDB Community Server & Mongoose ORM Docs. https://mongoosejs.com/docs/")
    add_bullet("Tailwind CSS Documentation. https://tailwindcss.com/")
    add_bullet("jsPDF Library Documentation. https://rawgit.com/MrRio/jsPDF/master/docs/")
    
    add_heading_2("9.3 GitHub Repositories")
    add_bullet("Bharath Rental System Source Code: https://github.com/shathanrajasekar-code/Car-Rental")

    doc.add_page_break()

    # ==========================================
    # CHAPTER 10: APPENDIX & TIME SHEET
    # ==========================================
    add_heading_1("CHAPTER 10: APPENDIX")
    add_para("", space_before=12)
    
    add_heading_2("10.1 Directory Folder Structure")
    dir_structure = (
        "car-rental-admin-system/\n"
        "├── client/                  # React Frontend App\n"
        "│   ├── src/                 # Component files, router, styles\n"
        "│   ├── index.html           # Main HTML page\n"
        "│   ├── package.json         # Frontend dependencies\n"
        "│   └── vite.config.js       # Vite build configurations\n"
        "├── server/                  # Express Backend Server\n"
        "│   ├── config/              # Database connection scripts\n"
        "│   ├── controllers/         # Express API controller handlers\n"
        "│   ├── models/              # Mongoose database schemas\n"
        "│   ├── routes/              # Express endpoint routing configuration\n"
        "│   ├── scripts/             # Database seeding scripts\n"
        "│   └── server.js            # Entry point for backend server\n"
        "├── car_rental_cli/          # Python command runner package\n"
        "│   └── main.py              # CLI executable script\n"
        "├── setup.py                 # Python package setup config\n"
        "├── pyproject.toml           # Python packaging settings\n"
        "├── run.bat                  # Double-click Windows launcher\n"
        "└── run.sh                   # Double-click macOS/Linux launcher\n"
    )
    add_para(dir_structure, italic=True, space_after=12)
    
    add_heading_2("10.2 Primary API Endpoints")
    api_list = [
        ("POST /api/auth/register", "Register a new customer or administrator"),
        ("POST /api/auth/login", "Authenticate credentials and return JWT session token"),
        ("GET /api/vehicles", "Fetch list of vehicles with dynamic brand-model filters"),
        ("POST /api/vehicles (Admin)", "Create a new vehicle entry in fleet list"),
        ("PUT /api/vehicles/:id (Admin)", "Update pricing, status, or details of a vehicle"),
        ("POST /api/bookings", "Validate dates, check overlaps, and create booking"),
        ("GET /api/bookings/my-bookings", "Retrieve history of bookings for current customer"),
        ("PUT /api/bookings/:id/status (Admin)", "Transition booking status (Ongoing, Completed, etc.)"),
        ("POST /api/reviews", "Post customer feedback ratings for a completed booking"),
        ("GET /api/admin/stats", "Aggregate analytics parameters for monthly income")
    ]
    
    t_api = doc.add_table(rows=1, cols=2)
    t_api.style = 'Table Grid'
    hdr = t_api.rows[0].cells
    hdr[0].text, hdr[1].text = "HTTP Route Endpoint", "Functional Description"
    hdr[0].paragraphs[0].runs[0].bold = True
    hdr[1].paragraphs[0].runs[0].bold = True
    for route, desc in api_list:
        row = t_api.add_row().cells
        row[0].text, row[1].text = route, desc
        for cell in row: set_cell_margins(cell, 60, 60, 100, 100)
        
    add_para("", space_before=12)
    
    add_heading_2("10.3 Project Timeline Time Sheet")
    add_para(
        "Below is the schedule tracking the weekly development stages of the Car Rental System:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    
    t_time = doc.add_table(rows=1, cols=5)
    t_time.style = 'Table Grid'
    hdr = t_time.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = "Week No", "Phase", "Key Tasks Performed", "Duration", "Status"
    for c in hdr: c.paragraphs[0].runs[0].bold = True
    
    timeline = [
        ("Week 1-2", "Analysis", "Gather requirements, study existing systems, draft specifications", "14 days", "Completed"),
        ("Week 3", "Design", "Draw Entity Relationship Diagrams (ERDs) and Data Flow Diagrams (DFDs)", "7 days", "Completed"),
        ("Week 4-5", "Database", "Design collections (User, Vehicle, Booking, Payment, Review) in MongoDB", "14 days", "Completed"),
        ("Week 6-7", "Backend", "Write Express API controllers, routing, JWT auth, and date-conflict engines", "14 days", "Completed"),
        ("Week 8-9", "Frontend", "Develop React components, cascading search catalogs, and billing cards", "14 days", "Completed"),
        ("Week 10", "Special Modules", "Program SafarLock warnings, FASTag toll simulator, and PDF invoicing", "7 days", "Completed"),
        ("Week 11", "Testing", "Perform unit tests, integration tests, and record test cases", "7 days", "Completed"),
        ("Week 12", "Deployment", "Create Python CLI setup, write run launchers, compile documentation", "7 days", "Completed")
    ]
    
    for wk, ph, tsk, dur, stat in timeline:
        row = t_time.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text, row[4].text = wk, ph, tsk, dur, stat
        for cell in row: set_cell_margins(cell, 60, 60, 100, 100)

    doc.add_page_break()

    # ==========================================
    # SCREENSHOTS PLACEHOLDERS
    # ==========================================
    add_heading_1("SCREENSHOTS & INTERFACES")
    add_para("", space_before=12)
    
    screens = [
        ("1. User Login Screen", "Allows users to log in securely using email and password.", "User inputs credentials -> Enforces required validations -> Returns JWT on success.", "JWT security, toggle visibility, custom warnings."),
        ("2. Vehicle Catalog & Cascading Filters", "Displays available cars based on brand-model selection.", "Select brand -> Model list filters -> Display matching vehicles.", "Cascading model dropdown, category tabs, availability filters."),
        ("3. Booking Confirmation Dialog", "Summarizes total costs including daily rental, deposit, and carbon offsetting.", "Pick dates -> Compute days and amount -> Check out with card mockup.", "Total price calculation, pre-paid toll options, carbon offset selection."),
        ("4. FASTag Toll Gates Simulator", "Simulates toll plaza crossings and balances dynamically.", "Click crossing plaza -> Deduct balance -> Update logs list.", "Plaza crossing triggers, live balance tracker, historical logs list."),
        ("5. SafarLock Speed Warning Dashboard", "Simulates speed checks and triggers violation alerts.", "Speed exceeds 120 km/h -> Red alerts -> Increment overspeed logs.", "Interactive speedometer, alert logging, violation counter."),
        ("6. Administrator Dashboard Charts", "Provides analytics on monthly revenues and bookings.", "Render charts based on bookings -> Show total count cards.", "Recharts bar charts, count status cards, active rentals grid."),
        ("7. PDF Invoice Download", "Generates client-side invoice logs in PDF format.", "Click download -> Client creates layout -> Save PDF file.", "Metadata table, price summary lists, status watermark.")
    ]
    
    for title, purp, desc, feat in screens:
        add_heading_2(title)
        add_para(f"**Purpose**: {purp}", space_after=2)
        add_para(f"**Description**: {desc}", space_after=2)
        add_para(f"**Key Features**: {feat}", space_after=4)
        
        # Screenshot Placeholder Box
        p_box = doc.add_paragraph()
        p_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_box.paragraph_format.space_before = Pt(12)
        p_box.paragraph_format.space_after = Pt(24)
        run_box = p_box.add_run("[ --- PLACEHOLDER: SCREENSHOT OF " + title.upper() + " --- ]")
        run_box.bold = True
        run_box.font.name = 'Times New Roman'
        run_box.font.size = Pt(11)
        run_box.font.color.rgb = RGBColor(128, 128, 128)

    # 3. Save Document
    doc_path = r"c:\Users\ragul\.gemini\antigravity-ide\scratch\car-rental-admin-system\Car_Rental_Administration_System_IET_Report.docx"
    doc.save(doc_path)
    print(f"IET Report successfully compiled and saved to: {doc_path}")

if __name__ == "__main__":
    create_document()
